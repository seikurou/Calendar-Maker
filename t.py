from docx import Document
from docxcompose.composer import Composer

document = Document('MonthlyCalendar.docx')
doc2 = Document('dd.docx')

composer = Composer(document)
composer.append(doc2)
composer.save("ss.docx")
# cell = document.tables[0].cell(1,0)
# cell.add_paragraph(text="hi", style=cell.paragraphs[0].style)
# cell.paragraphs[0].runs[0].clear().add_text("hi")
# document.add_table()
# document.save('dd.docx')