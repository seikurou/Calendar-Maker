from docx import Document
from datetime import date as Date
from calendar import Calendar, MonthlyCalendar

TEMPLATE = 'MonthlyLandscape.docx'
WEEK_LABELS = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']

class MonthlyLandscape(MonthlyCalendar):
    def __init__(self, date: Date, start_day: int=6):
        super().__init__(date, start_day)

    def makeDocument(self) -> Document:
        document = Document(TEMPLATE)
        self._labelWeeks(document)
        self._labelDays(document)
        return document

    
    def _labelWeeks(self, document: Document) -> None:
        start_day = self._start_day
        for col in range(7):
            cell = document.tables[0].cell(0, col)
            cell.paragraphs[0].runs[0].clear().add_text(WEEK_LABELS[start_day])
            start_day = (start_day + 1) % 7
    
    def _labelDays(self, document: Document) -> None:
        currDate = Date(self._date.year, self._date.month, 1)
        firstWeekday = currDate.weekday()
        month = currDate.month
        row = 1
        col = (firstWeekday - self._start_day) % 7
        while currDate.month == month:
            print(row, col)
            cell = document.tables[0].cell(row, col)
            cell.paragraphs[0].runs[0].clear().add_text(str(currDate.day))
            col += 1
            if col == 7:
                col = 0
                row += 2
            currDate = Calendar.nextDay(currDate)