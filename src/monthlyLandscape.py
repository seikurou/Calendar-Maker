from docx import Document
from datetime import date as Date
from src.calendar import Calendar, MonthlyCalendar

TEMPLATE = 'templates/MonthlyLandscape.docx'
WEEK_LABELS = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
MONTHS = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']

class MonthlyLandscape(MonthlyCalendar):
    def __init__(self, date: Date, start_day: int=6):
        super().__init__(date, start_day)

    def makeDocument(self) -> Document:
        document = Document(TEMPLATE)
        self._labelTitle(document)
        self._labelWeeks(document)
        self._labelDays(document)
        return document

    def _labelTitle(self, document: Document) -> None:
        document.paragraphs[0].runs[0].clear().add_text('{0:%B} {0:%Y}'.format(self._date))
    
    def _labelWeeks(self, document: Document) -> None:
        start_day = self._start_day
        for col in range(7):
            cell = document.tables[0].cell(0, col)
            cell.paragraphs[0].runs[0].clear().add_text(WEEK_LABELS[start_day])
            start_day = (start_day + 1) % 7
    
    def _labelDays(self, document: Document) -> None:
        currDate = Date(self._date.year, self._date.month, 1)
        firstWeekday = currDate.weekday() # weekday of the 1st. ex Jan 1 is on a Monday = (0)
        month = currDate.month
        startCol = (firstWeekday - self._start_day) % 7 # Starting column
        for r in range(1, 12, 2):
            for c in range(7):
                cell = document.tables[0].cell(r, c)
                cell = cell.paragraphs[0].runs[0].clear()
                if currDate.month == month and (r > 1 or c >= startCol):
                    cell.add_text(str(currDate.day))
                    currDate = Calendar.nextDay(currDate)
