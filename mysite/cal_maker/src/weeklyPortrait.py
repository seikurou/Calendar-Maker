from docx import Document
from datetime import date as Date
from .calendar import Calendar, WeeklyCalendar, TEMPLATE_PATH



class WeeklyPortrait(WeeklyCalendar):
    def __init__(self, date: Date, template: str, start_day: int=6, **kargs):
        super().__init__(date, template, start_day)

    def makeDocument(self) -> Document:
        document = Document(TEMPLATE_PATH + self._template)
        self._labelTitle(document)
        self._labelDays(document)
        return document

    def _labelTitle(self, document: Document) -> None:
        document.paragraphs[0].runs[0].clear().add_text('{0:%b} {0:%d}, {0:%Y}—{1:%b} {1:%d}, {1:%Y}'.format(self._date_start, self._date_end))

    def _labelDays(self, document: Document) -> None:
        currDate = self._date_start
        for r in range(7):
            cell = document.tables[0].cell(r, 0)
            cell = cell.paragraphs[0].runs[0].clear()
            cell.add_text('{0:%a}, {0:%b} {0:%d}'.format(currDate))
            currDate = Calendar.nextDay(currDate)
